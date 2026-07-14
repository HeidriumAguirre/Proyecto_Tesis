"""
Genera el Contrato de Confidencialidad en formato Word (.docx) para
el piloto del ITS PIE. Incluye:
- Encabezado con el logo UPA en la esquina superior derecha
- Texto del contrato en formato institucional
- Bloque de firmas al final

Si docs/assets/logo_upa.png no existe, usa un placeholder textual
"LOGO UPA" en lugar de la imagen.

Uso:
    python docs/build_contrato_docx.py
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
LOGO_PATH = ROOT / "docs" / "assets" / "logo_upa.png"
OUT_DOCX = ROOT / "docs" / "CONTRATO_CONFIDENCIALIDAD_DATOS.docx"

# Colores institucionales UPA (azules)
COLOR_UPA_AZUL = RGBColor(0x0F, 0x4C, 0x75)
COLOR_UPA_CELESTE = RGBColor(0x4A, 0x9D, 0xC7)
COLOR_GRIS = RGBColor(0x55, 0x55, 0x55)
COLOR_NEGRO = RGBColor(0x10, 0x10, 0x10)


def set_cell_border(cell, **kwargs):
    """Aplica bordes a una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = qn(f"w:{edge}")
            element = tcBorders.find(tag)
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tcBorders.append(element)
            element.set(qn("w:val"), kwargs[edge].get("val", "single"))
            element.set(qn("w:sz"), str(kwargs[edge].get("sz", 4)))
            element.set(qn("w:color"), kwargs[edge].get("color", "auto"))


def add_header_with_logo(doc):
    """Crea un encabezado con el logo UPA alineado a la derecha."""
    section = doc.sections[0]
    header = section.header

    # Limpiar parrafos existentes
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    # Crear una tabla invisible de 2 columnas: vacia a la izquierda, logo a la derecha
    table = header.add_table(rows=1, cols=2, width=Cm(17))
    table.autofit = False
    table.columns[0].width = Cm(10)
    table.columns[1].width = Cm(7)

    cell_left = table.rows[0].cells[0]
    cell_right = table.rows[0].cells[1]

    # Quitar bordes de la tabla del header
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                left={"val": "nil"},
                bottom={"val": "nil"},
                right={"val": "nil"},
            )

    # Celda izquierda: solo espacio
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Celda derecha: logo
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO_PATH.exists():
        run = p_right.add_run()
        try:
            run.add_picture(str(LOGO_PATH), width=Cm(5.5))
        except Exception as exc:
            p_right.add_run("[LOGO UPA - error al cargar: " + str(exc) + "]")
    else:
        run = p_right.add_run("LOGO UPA\n(guardar imagen en docs/assets/logo_upa.png)")
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_GRIS
        run.italic = True

    # Anadir un parrafo vacio debajo para evitar que el header quede muy pegado al body
    p_bottom = header.add_paragraph()
    p_bottom.paragraph_format.space_after = Pt(0)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COLOR_UPA_AZUL
    p.paragraph_format.space_after = Pt(6)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = COLOR_GRIS
    p.paragraph_format.space_after = Pt(12)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_UPA_AZUL
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_NEGRO
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_para(doc, text, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table_simple(doc, header, rows, col_widths_cm):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = False
    for i, w in enumerate(col_widths_cm):
        for cell in table.columns[i].cells:
            cell.width = Cm(w)
    # Header
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Background color
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "0F4C75")
        tcPr.append(shd)
    # Data
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    return table


def add_signature_block(doc, title, lines):
    """Crea un bloque de firma institucional con bordes."""
    add_h2(doc, title)
    table = doc.add_table(rows=len(lines), cols=2)
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(8.5)
        row.cells[1].width = Cm(8.5)
    for i, (label, value) in enumerate(lines):
        cell_l = table.rows[i].cells[0]
        cell_r = table.rows[i].cells[1]
        for cell, content, bold in [
            (cell_l, label, True),
            (cell_r, value, False),
        ]:
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(content)
            run.font.size = Pt(10)
            run.font.bold = bold
    # Espacio para firmar
    p_sign = doc.add_paragraph()
    p_sign.add_run("\n\n___________________________________").font.size = Pt(10)
    p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Firma y fecha")
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = COLOR_GRIS
    doc.add_paragraph()


def build_docx() -> None:
    doc = Document()

    # Estilos base: fuente Calibri 10
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Margenes
    section = doc.sections[0]
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Encabezado con logo
    add_header_with_logo(doc)

    # ========================================================================
    # PORTADA / TITULO
    # ========================================================================
    add_title(doc, "CONTRATO DE CONFIDENCIALIDAD Y AUTORIZACION")
    add_title(doc, "DE TRATAMIENTO DE DATOS")
    add_subtitle(
        doc,
        "Sistema de Tutoria Inteligente (ITS) basado en RAG para PIE — "
        "Colegio San Leonardo Murialdo de Valparaiso"
    )

    # ========================================================================
    # METADATOS
    # ========================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Version 1.0  -  Julio 2026")
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = COLOR_GRIS

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Tesista: Heidrium Nauj Aguirre Andrades  |  "
        "Profesor guia: Dr. Franklin Johnson  |  "
        "Universidad de Playa Ancha"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = COLOR_GRIS

    doc.add_paragraph()  # espacio

    # ========================================================================
    # 1. ANTECEDENTES
    # ========================================================================
    add_h1(doc, "1. Antecedentes y marco normativo")
    add_para(
        doc,
        "El presente contrato se suscribe en el marco del proyecto de tesis de pregrado "
        "senalado al inicio de este documento, y se rige por las siguientes "
        "disposiciones legales y reglamentarias chilenas:"
    )
    add_bullet(doc, "Ley N. 19.628 sobre Proteccion de la Vida Privada y Proteccion de Datos de Caracter Personal.")
    add_bullet(doc, "Decreto Supremo N. 170/2009 del Ministerio de Educacion, que fija normas para determinar los alumnos con necesidades educativas especiales que seran beneficiarios de la subvencion educacional especial.")
    add_bullet(doc, "Decreto Exento N. 83/2015 del Ministerio de Educacion, que aprueba criterios y orientaciones de adecuacion curricular para estudiantes con NEE.")
    add_bullet(doc, "Ley N. 20.422 sobre Igualdad de Oportunidades e Inclusion Social de Personas con Discapacidad.")
    add_bullet(doc, "Constitucion Politica de Chile, articulo 19 N. 4 (derecho a la vida privada) y N. 24 (derecho de propiedad sobre datos personales).")

    # ========================================================================
    # 2. OBJETO
    # ========================================================================
    add_h1(doc, "2. Objeto del contrato")
    add_para(
        doc,
        "Autorizar formalmente el tratamiento de datos anonimizados de 4 (cuatro) "
        "estudiantes del Programa de Integracion Escolar (PIE) del Colegio San "
        "Leonardo Murialdo de Valparaiso, con el unico fin de validar "
        "academicamente el prototipo de Sistema de Tutoria Inteligente en el marco "
        "del proyecto de tesis indicado al inicio de este documento."
    )

    # ========================================================================
    # 3. PARTES
    # ========================================================================
    add_h1(doc, "3. Partes")

    add_h2(doc, "3.1 Investigadora (responsable del tratamiento)")
    add_para(
        doc,
        "Heidrium Nauj Aguirre Andrades - RUT [RUT de la tesista] - "
        "Universidad de Playa Ancha, Facultad de Ingenieria, Escuela de Ingenieria "
        "Civil Informatica. Correo institucional: heidrium.aguirre@alumnos.upla.cl. "
        "Telefono: +56 X XXXX XXXX."
    )

    add_h2(doc, "3.2 Establecimiento educacional (custodio de los datos originales)")
    add_para(
        doc,
        "Colegio San Leonardo Murialdo de Valparaiso. Representante legal: "
        "[Nombre del sostenedor / representante legal] - RUT [RUT del colegio] - "
        "Domicilio: [Direccion del colegio] - Coordinadora PIE: Carolina Yanez "
        "(cyanez@murialdovalpo.cl)."
    )

    add_h2(doc, "3.3 Apoderado/a (autorizante del estudiante)")
    add_para(
        doc,
        "[Nombre del apoderado] - RUT [RUT del apoderado] - Domicilio: [Direccion "
        "del apoderado] - Correo: [Correo] - Telefono: [Telefono]. Estudiante "
        "asociado: [Nombre del estudiante - RUT anonimizado]."
    )

    # ========================================================================
    # 4. NATURALEZA Y ALCANCE DE LOS DATOS
    # ========================================================================
    add_h1(doc, "4. Naturaleza y alcance de los datos tratados")

    add_h2(doc, "4.1 Datos que se solicitan al colegio (anonimizados)")
    add_table_simple(
        doc,
        ["Categoria", "Campo", "Tratamiento"],
        [
            ["Identificacion", "RUT anonimizado (ficticio, formato XX.XXX.XXX-X)",
             "Reemplazo del RUT real por uno generado aleatoriamente"],
            ["Identificacion", "Nombre de pila (solo nombre y primer apellido)",
             "Seudonimizacion parcial"],
            ["Academica", "Curso y subdivision (A o B)", "Tal cual"],
            ["Academica", "Nivel de adaptacion de lenguaje (Alto, Medio, Bajo)", "Tal cual"],
            ["Pedagogica", "Requiere apoyo pictorico (Si/No)", "Tal cual"],
            ["Clinica (PIE)", "Codigos de diagnostico (TEA, TDAH, DIL, DEA, DL)", "Tal cual"],
        ],
        col_widths_cm=[3.0, 6.5, 7.5],
    )

    add_h2(doc, "4.2 Datos que NO se solicitaran bajo ninguna circunstancia")
    add_bullet(doc, "RUT real del estudiante.")
    add_bullet(doc, "Nombre completo del estudiante.")
    add_bullet(doc, "Direccion domiciliaria.")
    add_bullet(doc, "Datos de salud no contenidos en la categoria PIE (fichas clinicas, medicamentos, diagnosticos medicos no PIE).")
    add_bullet(doc, "Antecedentes familiares.")
    add_bullet(doc, "Datos de contacto del estudiante o su familia.")
    add_bullet(doc, "Fotografias, grabaciones de voz o video.")
    add_bullet(doc, "Informacion sobre rendimiento escolar especifico (calificaciones, asistencia).")
    add_bullet(doc, "Orientacion sexual, religion, origen etnico u otra informacion sensible.")

    add_h2(doc, "4.3 Datos generados durante el piloto")
    add_para(
        doc,
        "Durante el uso del prototipo se generara informacion de interaccion "
        "pedagogica que se almacenara en la base de datos del sistema: mensajes "
        "escritos o transcritos de audio intercambiados con el tutor socratico; "
        "fecha, hora y duracion de cada sesion; objetivo de aprendizaje (OA) "
        "consultado en cada sesion; y metricas agregadas (numero de sesiones, "
        "preguntas realizadas, duracion promedio)."
    )
    add_para(
        doc,
        "Estos datos se conservaran unicamente mientras dure el proyecto de tesis "
        "y se eliminaran al cierre del mismo (estimado marzo 2027), con excepcion "
        "de los resumenes estadisticos agregados que puedan incluirse en la memoria "
        "de titulo."
    )

    # ========================================================================
    # 5. FINALIDAD
    # ========================================================================
    add_h1(doc, "5. Finalidad del tratamiento")
    add_para(
        doc,
        "Los datos descritos en la seccion 4 se utilizaran exclusivamente para:"
    )
    add_bullet(doc, "Cargar perfiles PIE anonimizados en la base de datos del prototipo.")
    add_bullet(doc, "Personalizar las respuestas del tutor socratico segun el diagnostico y nivel del estudiante.")
    add_bullet(doc, "Validar la personalizacion del tutor mediante demostracion a docentes PIE (sin mostrar datos clinicos).")
    add_bullet(doc, "Analizar metricas agregadas de uso para la memoria de titulo de pregrado.")
    add_bullet(doc, "Publicar resultados agregados y anonimizados en la memoria de tesis y eventuales publicaciones academicas derivadas.")

    add_para(doc, "Queda expresamente prohibido:")
    add_bullet(doc, "Re-identificar a los estudiantes a partir de los datos anonimizados.")
    add_bullet(doc, "Compartir los datos con terceros no involucrados en el proyecto.")
    add_bullet(doc, "Usar los datos con fines comerciales, publicitarios o de vigilancia.")
    add_bullet(doc, "Difundir datos individuales (solo se publicaran metricas agregadas).")

    # ========================================================================
    # 6. MEDIDAS DE SEGURIDAD
    # ========================================================================
    add_h1(doc, "6. Medidas de seguridad y resguardo")
    add_para(
        doc,
        "La investigadora se compromete a aplicar las siguientes medidas:"
    )
    add_bullet(doc, "Anonimizacion previa de todos los datos antes de cargarlos al sistema.")
    add_bullet(doc, "Acceso restringido a la base de datos (solo la tesista y su profesor guia).")
    add_bullet(doc, "Contrasenas robustas (minimo 12 caracteres, alfanumericas, con simbolos) para todos los accesos.")
    add_bullet(doc, "No publicacion de datos identificables en la memoria de tesis, repositorios publicos (GitHub) o presentaciones.")
    add_bullet(doc, "Cifrado de la base de datos en transito (HTTPS) y en reposo (volumenes Docker cifrados a nivel de host).")
    add_bullet(doc, "Eliminacion definitiva de los datos al cierre del proyecto, salvo los resumenes estadisticos agregados.")
    add_bullet(doc, "Bitacora de auditoria de todas las acciones realizadas sobre los datos.")

    # ========================================================================
    # 7. COMPROMISOS DEL COLEGIO
    # ========================================================================
    add_h1(doc, "7. Compromisos del colegio")
    add_para(
        doc,
        "El Colegio San Leonardo Murialdo, a traves de su representante legal y "
        "de la Coordinadora PIE, se compromete a:"
    )
    add_bullet(doc, "Entregar los datos estrictamente anonimizados segun la plantilla PLANTILLA_INGESTA_ESTUDIANTES.csv proporcionada por la investigadora.")
    add_bullet(doc, "No entregar datos identificables reales bajo ninguna circunstancia.")
    add_bullet(doc, "Informar a los apoderados sobre el presente contrato y obtener su firma antes de la entrega de datos.")
    add_bullet(doc, "Custodiar el original firmado del contrato de cada apoderado.")
    add_bullet(doc, "Respetar el derecho del apoderado a retirar el consentimiento en cualquier momento, sin represalias para el estudiante.")
    add_bullet(doc, "Coordinar con la investigadora cualquier solicitud de acceso, rectificacion o eliminacion de datos.")

    # ========================================================================
    # 8. DERECHOS DEL APODERADO
    # ========================================================================
    add_h1(doc, "8. Derechos del apoderado/a y del estudiante")
    add_para(
        doc,
        "Conforme a la Ley N. 19.628, el apoderado y el estudiante tienen los "
        "siguientes derechos:"
    )
    add_bullet(doc, "Acceso: conocer que datos se tratan y para que fin.")
    add_bullet(doc, "Rectificacion: corregir datos inexactos.")
    add_bullet(doc, "Cancelacion: solicitar la eliminacion de los datos del sistema.")
    add_bullet(doc, "Oposicion: oponerse a un tratamiento especifico.")
    add_bullet(doc, "Revocacion del consentimiento: retirar el consentimiento otorgado en cualquier momento.")

    add_para(
        doc,
        "Estos derechos se ejercen mediante solicitud escrita dirigida a la "
        "investigadora al correo heidrium.aguirre@alumnos.upla.cl, quien debera "
        "responder en un plazo maximo de 10 dias habiles."
    )

    # ========================================================================
    # 9. DURACION Y REVOCACION
    # ========================================================================
    add_h1(doc, "9. Duracion y revocacion")
    add_para(
        doc,
        "El presente contrato tendra una duracion indefinida hasta el cierre del "
        "proyecto de tesis (estimado marzo 2027). El apoderado podra revocar el "
        "consentimiento en cualquier momento mediante comunicacion escrita a la "
        "investigadora o al colegio."
    )
    add_para(
        doc,
        "En caso de revocacion, los datos del estudiante seran eliminados en un "
        "plazo maximo de 30 dias desde la recepcion de la solicitud, manteniendo "
        "solo los resumenes estadisticos agregados ya generados (que no permiten "
        "re-identificacion)."
    )

    # ========================================================================
    # 10. CONFIDENCIALIDAD
    # ========================================================================
    add_h1(doc, "10. Confidencialidad de las partes")
    add_para(
        doc,
        "Las partes firmantes se obligan a mantener la mas estricta confidencialidad "
        "sobre: datos identificables de los estudiantes; diagnosticos clinicos "
        "especificos asociados a un estudiante; y cualquier informacion que pudiera "
        "permitir re-identificar a un estudiante. Esta obligacion se mantiene "
        "vigente aun despues de finalizado el proyecto y tiene caracter indefinido."
    )

    # ========================================================================
    # 11. SOLUCION DE CONTROVERSIAS
    # ========================================================================
    add_h1(doc, "11. Solucion de controversias")
    add_para(
        doc,
        "Cualquier controversia derivada del presente contrato sera resuelta "
        "preferentemente por acuerdo directo entre las partes. En caso de no "
        "alcanzarse acuerdo, las partes se someten a la competencia de los "
        "Tribunales Ordinarios de Justicia de Valparaisa, Chile."
    )

    # ========================================================================
    # 12. DISPOSICIONES FINALES
    # ========================================================================
    add_h1(doc, "12. Disposiciones finales")
    add_bullet(doc, "El presente contrato se firma en 3 (tres) ejemplares del mismo tenor, quedando uno en poder de cada parte (investigadora, colegio, apoderado).")
    add_bullet(doc, "La firma del apoderado es requisito habilitante para que el estudiante participe del piloto.")
    add_bullet(doc, "La firma del colegio formaliza la autorizacion institucional.")
    add_bullet(doc, "La firma de la investigadora declara el compromiso de cumplimiento.")

    # ========================================================================
    # FIRMAS
    # ========================================================================
    doc.add_page_break()
    add_h1(doc, "Firmas")

    add_h2(doc, "Investigadora")
    add_para(doc, "Nombre: Heidrium Nauj Aguirre Andrades")
    add_para(doc, "RUT: [RUT]")
    p = doc.add_paragraph()
    p.add_run("\n").font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("___________________________________")
    r.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Firma y fecha")
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = COLOR_GRIS

    doc.add_paragraph()
    add_h2(doc, "Representante del Colegio")
    add_para(doc, "Nombre: [Nombre del sostenedor / representante legal]")
    add_para(doc, "Cargo: [Cargo]")
    add_para(doc, "RUT: [RUT]")
    p = doc.add_paragraph()
    p.add_run("\n").font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("___________________________________")
    r.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Firma y fecha")
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = COLOR_GRIS

    doc.add_paragraph()
    add_h2(doc, "Coordinadora PIE (testigo tecnico)")
    add_para(doc, "Nombre: Carolina Yanez")
    add_para(doc, "Cargo: Coordinadora PIE - Colegio San Leonardo Murialdo")
    add_para(doc, "RUT: [RUT]")
    p = doc.add_paragraph()
    p.add_run("\n").font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("___________________________________")
    r.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Firma y fecha")
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = COLOR_GRIS

    doc.add_paragraph()
    add_h2(doc, "Apoderado/a")
    add_para(doc, "Nombre: [Nombre del apoderado]")
    add_para(doc, "RUT: [RUT]")
    add_para(doc, "Estudiante asociado: [Nombre del estudiante]")
    p = doc.add_paragraph()
    p.add_run("\n").font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("___________________________________")
    r.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Firma y fecha")
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = COLOR_GRIS

    # ========================================================================
    # ANEXO A
    # ========================================================================
    doc.add_page_break()
    add_h1(doc, "Anexo A - Resumen para el apoderado (1 carilla)")
    add_para(
        doc,
        "Este resumen puede entregarse impreso a cada apoderado en las reuniones "
        "de consentimiento. La firma se solicita sobre el contrato completo."
    )

    add_h2(doc, "CONSENTIMIENTO INFORMADO - Resumen")
    add_para(
        doc,
        "Su hijo/a es parte de un piloto educativo del Colegio San Leonardo "
        "Murialdo junto a la Universidad de Playa Ancha. En este piloto se prueba "
        "un tutor digital de matematicas que se adapta al diagnostico PIE de cada "
        "estudiante."
    )
    add_h2(doc, "Que datos entregaremos?")
    add_para(
        doc,
        "Solo el codigo del diagnostico PIE (por ejemplo, TEA, TDAH), el nivel "
        "de adaptacion de lenguaje y un RUT ficticio para identificar al alumno "
        "en el sistema. No entregaremos el RUT real, ni nombre completo, ni "
        "direccion, ni datos clinicos."
    )
    add_h2(doc, "Para que se usaran?")
    add_para(
        doc,
        "Para que el tutor digital entregue respuestas adaptadas a su hijo/a "
        "durante las sesiones de prueba, y para escribir la tesis de pregrado de "
        "la estudiante Heidrium Aguirre."
    )
    add_h2(doc, "Puedo retirarme?")
    add_para(
        doc,
        "Si, en cualquier momento y sin consecuencias para su hijo/a. Solo "
        "escribale a la investigadora."
    )
    add_h2(doc, "Quien vera los datos?")
    add_para(
        doc,
        "Solo la estudiante tesista y su profesor guia de universidad. No se "
        "compartiran con terceros."
    )
    add_para(
        doc,
        "Para mas detalles, lea el Contrato de Confidencialidad completo "
        "disponible con la Coordinadora PIE."
    )

    # Pie de pagina con version
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Contrato de Confidencialidad - Version 1.0 - Julio 2026")
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = COLOR_GRIS

    # Guardar
    doc.save(str(OUT_DOCX))
    print(f"Word generado: {OUT_DOCX}")
    print(f"Tamano: {OUT_DOCX.stat().st_size} bytes")
    if not LOGO_PATH.exists():
        print(f"")
        print(f"AVISO: Logo UPA no encontrado en {LOGO_PATH}")
        print(f"  -> Guarda la imagen del logo UPA como docs/assets/logo_upa.png")
        print(f"  -> Vuelve a ejecutar este script para regenerar el Word con el logo")


if __name__ == "__main__":
    build_docx()
