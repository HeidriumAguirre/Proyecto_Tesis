"""
Genera la Autorizacion de Tratamiento de Datos en formato Word (.docx).
Es el documento que firma el/la apoderado/a en las reuniones de
consentimiento. Es complementario al CONTRATO_CONFIDENCIALIDAD_DATOS.

Estructura:
- Encabezado con logo UPA (esquina superior derecha)
- Mismo formato de titulo que el contrato
- Un unico parrafo de autorizacion en primera persona con los
  datos de la tesista (Heidrium) y placeholders para los del
  colegio y el apoderado.
- 4 bloques de firma (Investigadora, Colegio, PIE, Apoderado).

Si docs/assets/logo_upa.png no existe, usa un placeholder textual.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
LOGO_PATH = ROOT / "docs" / "assets" / "logo_upa.png"
OUT_DOCX = ROOT / "docs" / "AUTORIZACION_TRATAMIENTO_DATOS.docx"

COLOR_UPA_AZUL = RGBColor(0x0F, 0x4C, 0x75)
COLOR_GRIS = RGBColor(0x55, 0x55, 0x55)
COLOR_NEGRO = RGBColor(0x10, 0x10, 0x10)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = qn(f"w:{edge}")
            element = tcBorders.find(tag)
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tcBorders.append(element)
            element.set(qn("w:val"), kwargs[edge].get("val", "single"))
            element.set(qn("w:sz"), str(kwargs[edge].get("sz", 4)))
            element.set(qn("w:color"), kwargs[edge].get("color", "auto"))


def add_header_with_logo(doc):
    section = doc.sections[0]
    header = section.header
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    table = header.add_table(rows=1, cols=2, width=Cm(17))
    table.autofit = False
    table.columns[0].width = Cm(10)
    table.columns[1].width = Cm(7)
    cell_left = table.rows[0].cells[0]
    cell_right = table.rows[0].cells[1]
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"val": "nil"}, left={"val": "nil"},
                           bottom={"val": "nil"}, right={"val": "nil"})
    cell_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO_PATH.exists():
        run = p_right.add_run()
        try:
            run.add_picture(str(LOGO_PATH), width=Cm(5.5))
        except Exception as exc:
            p_right.add_run("[LOGO UPA - error: " + str(exc) + "]")
    else:
        run = p_right.add_run("LOGO UPA\n(guardar imagen en docs/assets/logo_upa.png)")
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_GRIS
        run.italic = True
    header.add_paragraph()


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COLOR_UPA_AZUL
    p.paragraph_format.space_after = Pt(6)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = COLOR_GRIS
    p.paragraph_format.space_after = Pt(4)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_UPA_AZUL
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text, justify=True, first_line_indent_cm=0.0):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    return p


def add_signature_block(doc, role_data, role_title):
    """role_data = (label, value) por linea. role_title = 'Investigadora', etc."""
    add_h1(doc, role_title)
    table = doc.add_table(rows=len(role_data), cols=2)
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(13)
    for i, (label, value) in enumerate(role_data):
        cell_l = table.rows[i].cells[0]
        cell_r = table.rows[i].cells[1]
        cells_data = [
            (cell_l, label, True),
            (cell_r, value, False),
        ]
        for cell, txt, is_bold in cells_data:
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(10)
            run.font.bold = is_bold
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p_space = doc.add_paragraph()
    p_space.add_run("\n").font.size = Pt(10)
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_sign.add_run("___________________________________")
    r.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Firma y fecha")
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = COLOR_GRIS
    doc.add_paragraph()


def build_docx() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    section = doc.sections[0]
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    add_header_with_logo(doc)

    # ========================================================================
    # TITULO (mismo formato que el contrato)
    # ========================================================================
    add_title(doc, "CONTRATO DE CONFIDENCIALIDAD Y AUTORIZACION")
    add_title(doc, "DE TRATAMIENTO DE DATOS")
    add_subtitle(
        doc,
        "Sistema de Tutoria Inteligente (ITS) basado en RAG para PIE - "
        "Colegio San Leonardo Murialdo de Valparaiso"
    )
    add_subtitle(doc, "Version 1.0  -  Julio 2026")
    add_subtitle(
        doc,
        "Tesista: Heidrium Nauj Aguirre Andrades  |  "
        "Profesor guia: Dr. Franklin Johnson  |  "
        "Universidad de Playa Ancha"
    )

    doc.add_paragraph()

    # ========================================================================
    # AUTORIZACION (un solo parrafo en primera persona)
    # ========================================================================
    parrafo_autorizacion = (
        "Yo, (Nombre del apoderado/a - RUT), autorizo al Colegio San Leonardo Murialdo "
        "de Valparaiso, a su Coordinadora PIE (Datos de la persona) y a la estudiante de "
        "pregrado Heidrium Nauj Aguirre Andrades - RUT 21.063.023-6 - Correo institucional: "
        "heidrium.aguirre@alumnos.upla.cl - Telefono: +56 9 4202 3277 de la Universidad de "
        "Playa Ancha, Facultad de Ingenieria, Escuela de Ingenieria Civil Informatica, a "
        "utilizar los datos de mi alumno/a (Nombre/apellido/apodo - Curso - Subdivision - "
        "Nivel de adaptacion - Apoyo pedagogico - Diagnostico) con fines educativos para "
        "su tesis de pregrado."
    )
    add_para(doc, parrafo_autorizacion, justify=True, first_line_indent_cm=0.5)

    declaracion = (
        "Declaro haber leido y comprendido el Contrato de Confidencialidad y "
        "Autorizacion de Tratamiento de Datos (version 1.0, julio 2026), donde se "
        "detallan los datos solicitados, las medidas de seguridad, mis derechos como "
        "apoderado/a y la posibilidad de revocar este consentimiento en cualquier momento."
    )
    add_para(doc, declaracion, justify=True, first_line_indent_cm=0.5)

    # ========================================================================
    # FIRMAS
    # ========================================================================
    add_signature_block(
        doc,
        [
            ("Nombre", "Heidrium Nauj Aguirre Andrades"),
            ("RUT", "21.063.023-6"),
        ],
        "Investigadora",
    )

    add_signature_block(
        doc,
        [
            ("Nombre", "[Nombre del sostenedor / representante legal]"),
            ("Cargo", "[Cargo]"),
            ("RUT", "[RUT]"),
        ],
        "Representante del Colegio",
    )

    add_signature_block(
        doc,
        [
            ("Nombre", "Carolina Yanez"),
            ("Cargo", "Coordinadora PIE - Colegio San Leonardo Murialdo"),
            ("RUT", "[RUT]"),
        ],
        "Coordinadora PIE (testigo tecnico)",
    )

    add_signature_block(
        doc,
        [
            ("Nombre", "[Nombre del apoderado/a]"),
            ("RUT", "[RUT]"),
            ("Estudiante asociado", "[Nombre del estudiante]"),
        ],
        "Apoderado/a",
    )

    # Pie
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Autorizacion de Tratamiento de Datos - Version 1.0 - Julio 2026")
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = COLOR_GRIS

    doc.save(str(OUT_DOCX))
    print(f"Word generado: {OUT_DOCX}")
    print(f"Tamano: {OUT_DOCX.stat().st_size} bytes")
    if not LOGO_PATH.exists():
        print()
        print("AVISO: Logo UPA no encontrado en", LOGO_PATH)
        print("  -> Guarda la imagen del logo UPA como docs/assets/logo_upa.png")
        print("  -> Vuelve a ejecutar este script")


if __name__ == "__main__":
    build_docx()
